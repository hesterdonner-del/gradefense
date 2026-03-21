# Name: Wade:@
# Data: 2026/3/8 16:26

"""
主要目的：打开excel表格，进行计算，并生成三线表格

功能：
    1. 读取Excel文件（支持.xlsx, .xls）
    2. 可选数据处理：按指定列排序、添加计算列（如两列乘积）
    3. 生成Word文档，内含三线表（顶线、表头线、底线）

程序的基本流程：
    1.读取Excel文件（通过命令行参数或固定路径）。
    2.对数据进行处理（可选）。我们可以在代码中预设一些处理，比如按某列排序，或者添加汇总行。
    3.创建一个新的Word文档，添加一个标题（可选），然后插入表格。
    4.设置表格样式为三线表：
    通常第一行（表头）有上下边框，最后一行有下边框，表身没有竖线和横线。
    但在Word中实现三线表，我们通常需要：清除所有边框，然后为表格的第一行设置上边框和下边框（较粗？），为最后一行设置下边框。
    或者更符合标准三线表：顶线（表格上方），表头线（表头下方），底线（表格底部）。所以我们可以设置：
        表格无边框
        为第一行的上边框设置宽度
        为第一行的下边框设置宽度
        为最后一行的下边框设置宽度
        同时可能还需要设置表头行的下边框为双线或加粗？标准三线表通常顶线和底线较粗，表头线较细。我们可以根据需求调整。
    5.将DataFrame数据写入表格，注意表头和数据行。
    6.保存Word文档。

步骤：
    1.读取excel的数据
    2.计算数据
    3.读取word三线表模板
    4.将计算后的数据按照模板格式填入三线表
    5.保存新word

我们需要考虑以下步骤：
    读取Excel表格数据。可以使用pandas库，它非常方便。
    数据处理：
        可能包括计算、排序等。
        用户没有指定具体的处理逻辑，所以我们需要设计一个通用的程序，或者允许用户自定义一些操作。
        可以在程序中实现一些基本的数据处理选项，比如对某一列进行排序、计算两列的和或平均值等，或者简单允许用户通过参数指定。
        我们可以设计一个示例程序，包含一些简单的数据处理演示，比如读取后，对某列排序，或者添加一个计算列。
    生成三线表格并保存到Word文档。
        三线表格通常指只有三条横线的表格：顶线、表头线、底线。
        在Word中，我们可以使用python-docx库创建表格，并设置边框样式来实现三线表格。
        具体来说，可以创建一个表格，然后设置边框：只保留上边框、下边框和表头下边框。
        或者我们可以先设置表格无边框，然后手动给需要的行添加边框。

库：
    pandas      读取Excel和处理数据。
    python-docx 创建Word文档和表格。

参考：
    https://lbmd.seu.edu.cn/nano/info/1018/4207.htm
"""


import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import tkinter as tk
from tkinter import filedialog


def read_excel(file_path, sheet_name=0):
    """读取Excel文件，返回DataFrame"""
    try:
        df = pd.read_excel(file_path, sheet_name=sheet_name, dtype=str)  # 先按字符串读，避免日期等格式问题
        # 如果希望保留原始数据类型，可去掉dtype参数，但日期列在后续写入Word时会变成datetime对象，可能需要转换
        return df
    except Exception as e:
        print(f"读取Excel失败：{e}")
        return None


def process_data(df, sort_column=None, ascending=True, calc_columns=None):
    """
    数据处理函数
    :param df: 原始DataFrame
    :param sort_column: 按哪一列排序，None表示不排序
    :param ascending: 排序顺序
    :param calc_columns: 计算列定义列表，每个元素为(新列名, 表达式字符串)，表达式中的列名用df['列名']表示
                         例如：[('总价', 'df["数量"] * df["单价"]')]
    :return: 处理后的DataFrame
    """
    # 深拷贝避免修改原数据
    df = df.copy()

    # 执行计算列
    if calc_columns:
        for new_col, expr in calc_columns:
            try:
                # 注意：eval有安全风险，但此处假设用户输入可信；生产环境建议使用pandas内置运算
                df[new_col] = eval(expr)
            except Exception as e:
                print(f"计算列 {new_col} 失败：{e}")

    # 排序
    if sort_column and sort_column in df.columns:
        df = df.sort_values(by=sort_column, ascending=ascending).reset_index(drop=True)

    return df


def set_cell_border(cell, border_type="top", sz=1, color="000000"):
    """
    设置单元格边框
    :param cell: 单元格对象
    :param border_type: 边框类型，如"top","bottom","left","right"
    :param sz: 边框宽度（1/8磅）
    :param color: RGB颜色字符串，如"000000"
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tcPr.append(borders)

    border = borders.find(qn(f'w:{border_type}'))
    if border is None:
        border = OxmlElement(f'w:{border_type}')
        borders.append(border)

    border.set(qn('w:val'), 'single')          # 单线
    border.set(qn('w:sz'), str(sz))            # 线宽
    border.set(qn('w:color'), color)           # 颜色


def apply_three_line_style(table):
    """
    将表格应用三线表样式：
    - 顶线：第一行的上边框（加粗）
    - 表头线：第一行的下边框（普通）
    - 底线：最后一行的下边框（加粗）
    同时清除所有单元格的其他边框（左、右、内部横线）
    """
    # 先清除所有边框（所有单元格的所有边框）
    for row in table.rows:
        for cell in row.cells:
            # 移除所有边框（通过设置边框为None或者不设置；这里我们显式移除所有边框元素）
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is not None:
                borders = tcPr.find(qn('w:tcBorders'))
                if borders is not None:
                    tcPr.remove(borders)

    # 设置顶线（第一行每个单元格的上边框）
    first_row = table.rows[0]
    for cell in first_row.cells:
        set_cell_border(cell, border_type="top", sz=12)  # 顶线加粗，24*1/8=1.5磅

    # 设置表头线（第一行每个单元格的下边框）
    for cell in first_row.cells:
        set_cell_border(cell, border_type="bottom", sz=6)  # 表头线0.75磅

    # 设置底线（最后一行每个单元格的下边框）
    last_row = table.rows[-1]
    for cell in last_row.cells:
        set_cell_border(cell, border_type="bottom", sz=12)  # 底线加粗1.5磅


def dataframe_to_word_table(doc, df, title=None):
    """
    将DataFrame写入Word文档，生成三线表
    :param doc: Document对象
    :param df: DataFrame
    :param title: 表格标题（可选）
    """
    if title:
        doc.add_heading(title, level=2)

    # 创建表格：行数=数据行数+表头，列数=df列数
    table = doc.add_table(rows=len(df)+1, cols=len(df.columns))
    table.style = 'Table Grid'  # 先使用网格样式，之后会覆盖边框
    table.autofit = True
    table.allow_autofit = True

    # 写入表头
    for j, col_name in enumerate(df.columns):
        cell = table.cell(0, j)
        cell.text = str(col_name)
        # 表头加粗
        cell.paragraphs[0].runs[0].font.bold = True

    # 写入数据行
    for i in range(len(df)):
        for j, col_name in enumerate(df.columns):
            value = df.iloc[i, j]
            # 处理NaN或None
            if pd.isna(value):
                value = ""
            cell = table.cell(i+1, j)
            cell.text = str(value)

    # 应用三线表样式
    apply_three_line_style(table)

    # 可选：设置单元格垂直居中，字体大小等
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_TABLE_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = '宋体'
                    # 中文字体名称需要设置东亚字体
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def open_file_or_folder():
    root = tk.Tk()
    root.withdraw()

    FolderName = filedialog.askdirectory()  # 获取文件夹
    FileName = filedialog.askopenfilename()  # 获取文件夹中的某文件

    if '/' in FolderName:
        FolderName.replace('/', '\\')
        print(FolderName, 'Finding 1')

    if '/' in FileName:
        FileName.replace('/', '\\')
        print(FileName, 'Finding 2')

    if len(FolderName) == 0:
        print('No Folder Name Found!')
    else:
        print('Folder Name Found! And Folder is:', FolderName)

    if len(FileName) == 0:
        print('No File Name Found!')
    else:
        print('File Name Found! And File is:', FileName)

    return FileName , FolderName

def main():
    # ==================== 配置区域 ====================
    input_excel = "D:\\wlg2o\\Documents\\景艺大\\课程\\25_1_耐火材料工艺学-董明堂\\作业_三线表格.xlsx"          # 输入Excel文件路径
    output_word = "D:\\wlg2o\\Documents\\景艺大\\课程\\25_1_耐火材料工艺学-董明堂\\output.docx"        # 输出Word文件路径
    sheet = 0                           # 工作表索引或名称
    sort_column = "B"                 # 按此列排序，若无需排序设为None
    ascending = False                   # 降序
    # 计算列定义：每个元组 (新列名, 表达式)，表达式可用df['列名']引用其他列
    calc_columns = [
        ("b", 'df["1"] * df["2"]')
    ]
    # =================================================

    # 1. 读取Excel
    df = read_excel(input_excel, sheet)
    if df is None:
        return

    print("原始数据预览：", df.head())

    # 2. 数据处理（排序、计算列）
    df_processed = process_data(df, sort_column=sort_column, ascending=ascending, calc_columns=calc_columns)

    print("处理后数据预览：", df_processed.head())

    # 3. 创建Word文档
    doc = Document()

    # 设置文档默认字体（可选）
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 添加标题
    # doc.add_heading('数据三线表', level=1)

    # 4. 将DataFrame写入表格
    dataframe_to_word_table(doc, df_processed, title="统计结果")

    # 5. 保存文档
    doc.save(output_word)
    print(f"Word文档已生成：{output_word}")


if __name__ == "__main__":
    main()
