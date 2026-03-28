# Name: Wade:@
# Data: 2026/3/22 19:31

import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.shared import Pt, Cm, RGBColor
import tkinter as tk
from tkinter import filedialog
import openpyxl


def main():
    # 选择文件夹
    file_name, folder_name = open_file_folder()
    # print(file_name)

    print_excel(file_name)
    sheet_name = input("请输入此文件需要处理的页：")

    # 读取Excel
    df = read_excel(file_path=file_name, sheet_name=sheet_name)
    # print(df.head())

    doc = Document()
    # 设置文档默认字体（可选）
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    dataframe_to_tlt(df, doc)
    doc.save(".\\appendix\\text.docx")


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            element = tcBorders.find(qn('w:' + edge))
            if element is None:
                element = OxmlElement('w:' + edge)
                tcBorders.append(element)

            for key in ["sz", "val", "color", "space"]:
                if key in edge_data:
                    element.set(qn('w:' + key), str(edge_data[key]))

def remove_all_borders(table):
    # 清除所有单元格的边框
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "nil"},
                bottom={"val": "nil"},
                left={"val": "nil"},
                right={"val": "nil"}
            )

def apply_tlt(table):
    """
    设置单元格边框
    cell: 单元格对象
    border_type: 边框类型，如"top","bottom","left","right"
    sz: 边框宽度（1/8磅）
    color: RGB颜色字符串，如"000000"
    """

    for cell in table.rows[0].cells:    # 设置顶线
        set_cell_border(cell, top={"sz": 12, "val": "single", "color": "000000"})

    for cell in table.rows[0].cells:    # 设置表头线
        set_cell_border(cell, bottom={"sz": 6, "val": "single", "color": "000000"})

    for cell in table.rows[-1].cells:   # 设置底线
        set_cell_border(cell, bottom={"sz": 12, "val": "single", "color": "000000"})


def open_file_folder():
    root = tk.Tk()
    root.withdraw()

    FolderName = filedialog.askdirectory()  # 获取文件夹
    FileName = filedialog.askopenfilename()  # 获取文件夹中的某文件

    if '/' in FolderName:
        FolderName.replace('/', '\\')
        print(FolderName, 'Finding 1')

    if '/' in FileName:
        FileName.replace('/', '\\')
        print(FileName, 'Finding 2')

    if len(FolderName) == 0:
        print('未找到文件夹!')
    else:
        print('文件夹找到了! 文件夹地址:', FolderName)

    if len(FileName) == 0:
        print('未找到文件!')
    else:
        print('文件找到了! 文件地址:', FileName)

    return FileName , FolderName

def read_excel(file_path, sheet_name):
    # 读取Excel文件，返回 DataFrame
    try:
        df = pd.read_excel(file_path, sheet_name=sheet_name, dtype=str)  # 先按字符串读，避免日期等格式问题
        # 如果希望保留原始数据类型，可去掉dtype参数，但日期列在后续写入Word时会变成datetime对象，可能需要转换
        return df
    except Exception as e:
        print(f"读取Excel失败：{e}")
        return None

def print_excel(file_path):
    all_sheets = pd.ExcelFile(file_path)
    print(all_sheets.sheet_names)
    return None

def dataframe_to_tlt(df, doc, caption=None, footnote=None):
    # 将 DataFrame 数据转换为 Word 表格
    rows, cols = df.shape

    table = doc.add_table(rows=rows + 1, cols=cols)

    for j, col in enumerate(df.columns):
        cell = table.cell(0, j)
        cell.text = col
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i in range(rows):
        for j in range(cols):
            values = df.iloc[i, j]
            cell = table.cell(i+1, j)
            cell.text = str(values)

            # 自动右对齐数字
            if isinstance(values, (int, float)):
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    remove_all_borders(table)
    apply_tlt(table)

    return table


if __name__ == '__main__':
   main()
